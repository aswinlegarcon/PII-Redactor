from flask import Flask, request, jsonify, send_file
from flask_cors import CORS
from PIL import Image
import pytesseract
import pdfplumber
import os
import io
import re
from docx import Document
import cv2
import numpy as np
from presidio_analyzer import AnalyzerEngine, RecognizerResult
from presidio_anonymizer import AnonymizerEngine
from reportlab.lib.pagesizes import letter
from reportlab.pdfgen import canvas

app = Flask(__name__)
CORS(app)

# Initialize Presidio engines
analyzer = AnalyzerEngine()
anonymizer = AnonymizerEngine()

# Predefined coordinates for Aadhaar elements
COORDINATES = {
    'AADHAAR_NUMBER': (336, 473, 675, 536),
    'NAME': (312, 145, 525, 228),
    'DOB': (316, 237, 662, 276)
}

def blur_region(image, coordinates):
    """Blur a region in an image based on the provided coordinates."""
    x1, y1, x2, y2 = coordinates
    roi = image[y1:y2, x1:x2]
    roi = cv2.GaussianBlur(roi, (51, 51), 0)
    image[y1:y2, x1:x2] = roi
    return image

def process_file(file):
    """Process the uploaded file and return the text content."""
    file_ext = os.path.splitext(file.filename)[1].lower()
    file_buffer = io.BytesIO(file.read())
    file_buffer.seek(0)

    try:
        if file_ext == '.txt':
            return file_buffer.getvalue().decode('utf-8'), 'txt'
        elif file_ext in ['.png', '.jpg', '.jpeg']:
            image = Image.open(file_buffer)
            text = pytesseract.image_to_string(image)
            return text, 'txt'
        elif file_ext == '.pdf':
            text = ""
            try:
                with pdfplumber.open(file_buffer) as pdf:
                    for page in pdf.pages:
                        extracted_text = page.extract_text()
                        if extracted_text:
                            text += extracted_text + "\n"
            except Exception as e:
                print(f"Error processing PDF: {e}")
                return None, None
            return text, 'txt'
        elif file_ext == '.docx':
            try:
                doc = Document(file_buffer)
                text = ""
                for paragraph in doc.paragraphs:
                    text += paragraph.text + "\n"
                return text, 'docx'
            except Exception as e:
                print(f"Error processing DOCX: {e}")
                return None, None
    except Exception as e:
        print(f"Error in process_file: {e}")
        return None, None


def remove_pii(text):
    """Analyze and remove PII from the given text."""
    results = analyzer.analyze(text=text, entities=[], language='en')
    if results:
        pii_detected = True
        redacted_text = anonymizer.anonymize(text=text, analyzer_results=results).text
        return redacted_text, pii_detected
    else:
        return text, False

def save_processed_file(text, file_ext):
    """Save the processed text into a file and return the file path."""
    output_buffer = io.BytesIO()

    if file_ext == 'txt':
        c = canvas.Canvas(output_buffer, pagesize=letter)
        text_object = c.beginText(40, 750)
        for line in text.splitlines():
            text_object.textLine(line)
        c.drawText(text_object)
        c.save()
        output_buffer.seek(0)
        return output_buffer, 'processed.pdf'
    elif file_ext == 'docx':
        doc = Document()
        for line in text.splitlines():
            doc.add_paragraph(line)
        doc.save(output_buffer)
        output_buffer.seek(0)
        return output_buffer, 'processed.docx'

    return None, None

@app.route('/upload', methods=['POST'])
def upload_file():
    if 'file' not in request.files:
        return jsonify({'error': 'No file part'}), 400
    
    file = request.files['file']
    if file.filename == '':
        return jsonify({'error': 'No selected file'}), 400

    file_ext = os.path.splitext(file.filename)[1].lower()

    if file_ext in ['.png', '.jpg', '.jpeg']:
        # Convert the uploaded image to OpenCV format
        file_buffer = np.frombuffer(file.read(), np.uint8)
        image = cv2.imdecode(file_buffer, cv2.IMREAD_COLOR)

        # Get user selection from form data
        blur_aadhaar = request.form.get('blur_aadhaar') == 'true'
        blur_name = request.form.get('blur_name') == 'true'
        blur_dob = request.form.get('blur_dob') == 'true'

        # Apply blurring based on user selection
        if blur_aadhaar:
            image = blur_region(image, COORDINATES['AADHAAR_NUMBER'])
        if blur_name:
            image = blur_region(image, COORDINATES['NAME'])
        if blur_dob:
            image = blur_region(image, COORDINATES['DOB'])

        # Convert the processed image back to PIL format
        _, buffer = cv2.imencode('.png', image)
        image_io = io.BytesIO(buffer)
        return send_file(image_io, mimetype='image/png', as_attachment=True, download_name='processed_image.png')
    else:
        # Process text-based files
        file.seek(0)
        text, file_ext = process_file(file)

        if text:
            redacted_text, pii_detected = remove_pii(text)
            output_buffer, filename = save_processed_file(redacted_text, file_ext)

            if output_buffer:
                # Set MIME type based on the file extension
                mimetype = 'application/vnd.openxmlformats-officedocument.wordprocessingml.document' if file_ext == 'docx' else 'application/pdf'
                return send_file(output_buffer, as_attachment=True, download_name=filename, mimetype=mimetype), 200, {'PII-Detected': 'true' if pii_detected else 'false'}

    return jsonify({'error': 'Unable to process the file.'}), 500

if __name__ == '__main__':
    app.run(debug=True)

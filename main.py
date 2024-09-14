import cv2
import sys
import os
import io
import re
from PIL import Image
import pytesseract
import pdfplumber
from docx import Document
from presidio_analyzer import AnalyzerEngine, RecognizerResult
from presidio_anonymizer import AnonymizerEngine
from fpdf import FPDF


# Initialize Presidio engines
analyzer = AnalyzerEngine()
anonymizer = AnonymizerEngine()

# Predefined coordinates for Aadhaar elements
COORDINATES = {
    'AADHAAR_NUMBER': (336, 473, 675, 536),
    'NAME': (312, 145, 525, 228),
    'DOB': (316, 237, 662, 276)
}

def blur_region(image, coordinates):
    """Blur a region in an image based on the provided coordinates."""
    x1, y1, x2, y2 = coordinates
    roi = image[y1:y2, x1:x2]
    roi = cv2.GaussianBlur(roi, (51, 51), 0)
    image[y1:y2, x1:x2] = roi
    return image

def process_image(file_path, blur_aadhaar, blur_name, blur_dob):
    """Process the image to blur selected PII regions and save as PNG."""
    image = cv2.imread(file_path)

    if image is None:
        raise Exception(f"Failed to load image: {file_path}")

    if blur_aadhaar == 1:
        image = blur_region(image, COORDINATES['AADHAAR_NUMBER'])
    if blur_name == 1:
        image = blur_region(image, COORDINATES['NAME'])
    if blur_dob == 1:
        image = blur_region(image, COORDINATES['DOB'])

    # Save the image as PNG with "processed_" prefix
    input_filename = os.path.splitext(os.path.basename(file_path))[0]
    output_file = f"processed_{input_filename}.png"
    output_file_path = os.path.join(os.getcwd(), output_file)

    cv2.imwrite(output_file_path, image, [cv2.IMWRITE_PNG_COMPRESSION, 9])

    print(f"Processed image saved at: {output_file_path}")
    return output_file_path


def process_text(text):
    """Remove PII from the text using Presidio."""
    results = analyzer.analyze(text=text, language='en')

    # Convert results to a format that can be used by the anonymizer
    anonymize_results = [RecognizerResult(
        entity_type=result.entity_type,
        start=result.start,
        end=result.end,
        score=result.score
    ) for result in results]

    # Regex for Aadhaar numbers
    aadhaar_regex = r'\b\d{4} \d{4} \d{4}\b'
    for match in re.finditer(aadhaar_regex, text):
        anonymize_results.append(RecognizerResult(
            entity_type="AADHAAR_NUMBER",
            start=match.start(),
            end=match.end(),
            score=1.0
        ))

    redacted_text = text

    # Anonymize the text based on the detected entities
    if anonymize_results:
        redacted_text = anonymizer.anonymize(text=redacted_text, analyzer_results=anonymize_results).text

    # Additional replacement for Aadhaar numbers
    redacted_text = re.sub(aadhaar_regex, '<REDACTED AADHAAR>', redacted_text)

    return redacted_text

def save_processed_file(text, file_path, output_format):
    """Save the processed text into a file with the same naming convention."""
    input_filename = os.path.splitext(os.path.basename(file_path))[0]
    
    if output_format == 'pdf':
        pdf = FPDF()
        pdf.add_page()
        pdf.set_auto_page_break(auto=True, margin=15)
        pdf.set_font("Arial", size=12)
        for line in text.splitlines():
            pdf.multi_cell(0, 10, line)
        output_file = f'processed_{input_filename}.pdf'
        pdf.output(output_file)
        print(f"Processed PDF saved at: {output_file}")
        return output_file

    elif output_format == 'docx':
        doc = Document()
        for line in text.splitlines():
            doc.add_paragraph(line)
        output_file = f'processed_{input_filename}.docx'
        doc.save(output_file)
        print(f"Processed DOCX saved at: {output_file}")
        return output_file

    return None

def process_file(file_path):
    """Process the uploaded file and return the text content."""
    file_ext = os.path.splitext(file_path)[1].lower()

    if file_ext == '.txt':
        with open(file_path, 'r', encoding='utf-8') as f:
            return f.read(), 'txt'
    elif file_ext in ['.png', '.jpg', '.jpeg']:
        image = Image.open(file_path)
        text = pytesseract.image_to_string(image)
        return text, 'txt'
    elif file_ext == '.pdf':
        text = ""
        with pdfplumber.open(file_path) as pdf:
            for page in pdf.pages:
                extracted_text = page.extract_text()
                if extracted_text:
                    text += extracted_text + "\n"
        return text, 'pdf'
    elif file_ext == '.docx':
        doc = Document(file_path)
        text = ""
        for paragraph in doc.paragraphs:
            text += paragraph.text + "\n"
        return text, 'docx'
    else:
        raise ValueError(f"Unsupported file format: {file_ext}")

def main():
    """Main function to process files."""
    if len(sys.argv) < 2:
        print("Usage: python3 script.py <file_path> [blur_aadhaar] [blur_name] [blur_dob]")
        return

    file_path = sys.argv[1]
    file_ext = os.path.splitext(file_path)[1].lower()

    if file_ext in ['.png', '.jpg', '.jpeg']:
        # Process the image
        blur_aadhaar = int(sys.argv[2]) if len(sys.argv) > 2 else 0
        blur_name = int(sys.argv[3]) if len(sys.argv) > 3 else 0
        blur_dob = int(sys.argv[4]) if len(sys.argv) > 4 else 0

        output_file = process_image(file_path, blur_aadhaar, blur_name, blur_dob)
        print(f"Processed image saved as: {output_file}")

    elif file_ext in ['.pdf', '.docx', '.txt']:
        # Process text-based files
        text, output_format = process_file(file_path)
        redacted_text = process_text(text)
        output_file = save_processed_file(redacted_text, file_path, output_format)
        print(f"Processed text saved as: {output_file}")
    else:
        print(f"Unsupported file format: {file_ext}")

if __name__ == '__main__':
    main()
